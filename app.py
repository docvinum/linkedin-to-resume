import os
import sys
from PyPDF2 import PdfReader
import markdown
from docx import Document

# Fonction pour vérifier la présence du fichier PDF
def check_pdf_file(file_name="Profile.pdf"):
    if not os.path.isfile(file_name):
        print(f"Erreur : Le fichier {file_name} est introuvable dans le dossier actuel.")
        sys.exit(1)
    return file_name

# Fonction pour extraire les données du PDF
def extract_data_from_pdf(file_path):
    try:
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    except Exception as e:
        print(f"Erreur lors de l'extraction des données : {e}")
        sys.exit(1)

# Fonction pour formater les données extraites
def format_data(raw_text):
    # Exemple de traitement pour structurer les informations
    data = {
        "contact": extract_contact_info(raw_text),
        "skills": extract_skills(raw_text),
        "languages": extract_languages(raw_text),
        "certifications": extract_certifications(raw_text),
        "experience": extract_experience(raw_text),
        "education": extract_education(raw_text),
        "publications": extract_publications(raw_text),
    }
    return data

def extract_contact_info(text):
    # Extraction basique pour email et LinkedIn (à améliorer avec regex)
    email = "exemple@mail.com"  # Ajouter une extraction réelle
    linkedin = "linkedin.com/exemple"  # Ajouter une extraction réelle
    return {"email": email, "linkedin": linkedin}

# Ajoutez les autres fonctions d'extraction ici...

# Fonction pour exporter les données dans différents formats
def export_to_format(data, export_format="md"):
    if export_format == "md":
        file_name = "profil.md"
        with open(file_name, "w") as f:
            f.write(markdown.markdown(str(data)))
    elif export_format == "txt":
        file_name = "profil.txt"
        with open(file_name, "w") as f:
            f.write(str(data))
    elif export_format == "docx":
        file_name = "profil.docx"
        doc = Document()
        doc.add_heading("CV Généré", level=1)
        for key, value in data.items():
            doc.add_heading(key.capitalize(), level=2)
            doc.add_paragraph(str(value))
        doc.save(file_name)
    else:
        print("Erreur : Format non pris en charge.")
        sys.exit(1)
    print(f"Le CV a été généré avec succès sous le format : {export_format}")
    return file_name

# Fonction principale
def main():
    pdf_file = check_pdf_file()
    raw_text = extract_data_from_pdf(pdf_file)
    structured_data = format_data(raw_text)
    
    print("Choisissez le format d'export :")
    print("1. Markdown (.md)")
    print("2. Texte brut (.txt)")
    print("3. Word (.docx)")
    choice = input("Entrez le numéro de votre choix : ")

    if choice == "1":
        export_to_format(structured_data, "md")
    elif choice == "2":
        export_to_format(structured_data, "txt")
    elif choice == "3":
        export_to_format(structured_data, "docx")
    else:
        print("Choix invalide.")

if __name__ == "__main__":
    main()
